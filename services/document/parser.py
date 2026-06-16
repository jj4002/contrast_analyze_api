def parse_docx(file_path: str) -> str:
    from docx import Document as DocxDocument
    doc = DocxDocument(file_path)
    text_parts = []
    for p in doc.paragraphs:
        t = p.text.strip()
        if t:
            text_parts.append(t)
    for table in doc.tables:
        for row in table.rows:
            rt = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
            if rt:
                text_parts.append(rt)
    if not text_parts:
        raise ValueError("No text could be extracted from the DOCX file")
    return "\n".join(text_parts)


def parse_pdf(file_path: str) -> str:
    import pdfplumber
    text_parts = []
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            t = page.extract_text()
            if t and t.strip():
                text_parts.append(t)
    if not text_parts:
        raise ValueError("No text could be extracted from the PDF file")
    return "\n".join(text_parts)


def parse_document(file_path: str, file_ext: str) -> str:
    ext = file_ext.lower()
    if ext in (".docx", ".doc"):
        return parse_docx(file_path)
    elif ext == ".pdf":
        return parse_pdf(file_path)
    raise ValueError(f"Unsupported file format: {ext}")
